"""
Knowledge Base service for customer document management and ICP extraction

Handles document upload, parsing, embedding generation, and vector similarity search
Uses RunPod S3 storage and PostgreSQL for metadata
"""
import os
import io
import hashlib
from typing import List, Dict, Optional
from datetime import datetime

# Document parsing
# PDF processing - optional dependency
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

# DOCX processing - optional dependency
try:
    from docx import Document as DocxDocument
except ImportError:
    DocxDocument = None

# Embeddings - optional dependency
try:
    from sentence_transformers import SentenceTransformer
    import numpy as np
    SENTENCE_TRANSFORMERS_AVAILABLE = True
except ImportError:
    SentenceTransformer = None
    np = None
    SENTENCE_TRANSFORMERS_AVAILABLE = False

# RunPod storage and database
from sqlalchemy.orm import Session
from app.services.runpod_storage import RunPodStorageService
from app.models.customer_models import KnowledgeDocument
from app.core.logging import setup_logging

logger = setup_logging(__name__)


class KnowledgeBaseService:
    """
    Service for customer knowledge base management

    Features:
    - Document upload (PDF, DOCX, TXT) to RunPod S3 Storage
    - Text extraction and chunking
    - Vector embedding generation for similarity search
    - ICP (Ideal Customer Profile) extraction
    - Multi-tenant document isolation
    - PostgreSQL metadata storage
    """

    def __init__(self):
        """Initialize Knowledge Base service with RunPod storage and embedding model"""
        # Initialize RunPod S3 storage
        self.storage = RunPodStorageService(datacenter="US-CA-2")

        # Initialize embedding model (384-dimensional embeddings)
        # Using all-MiniLM-L6-v2 for fast, efficient embeddings
        self.embedding_model_name = os.getenv(
            "EMBEDDING_MODEL",
            "sentence-transformers/all-MiniLM-L6-v2"
        )
        self.embedding_model = SentenceTransformer(self.embedding_model_name)
        self.embedding_dimension = 384  # Dimension for all-MiniLM-L6-v2

        logger.info(f"Initialized Knowledge Base with embedding model: {self.embedding_model_name}")

    def upload_document(
        self,
        file_content: bytes,
        filename: str,
        customer_id: str,
        content_type: str,
        db: Session
    ) -> Dict:
        """
        Upload and process a customer document

        Args:
            file_content: Document file bytes
            filename: Original filename
            customer_id: Customer ID for isolation
            content_type: MIME type
            db: Database session

        Returns:
            Dictionary with document metadata and extracted content
        """
        try:
            # Generate unique document ID
            doc_hash = hashlib.sha256(file_content).hexdigest()[:16]
            timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
            doc_id = f"{customer_id}_{timestamp}_{doc_hash}"

            # Upload to RunPod S3 Storage
            storage_path = f"customers/{customer_id}/documents/{doc_id}/{filename}"
            file_obj = io.BytesIO(file_content)
            file_url = self.storage.upload_fileobj(
                file_obj=file_obj,
                object_name=storage_path,
                content_type=content_type
            )

            # Extract text from document
            extracted_text = self._extract_text(file_content, filename, content_type)

            # Generate embedding for the full document
            embedding = self._generate_embedding(extracted_text)

            # Extract ICP criteria from document
            icp_data = self._extract_icp_criteria(extracted_text)

            # Store document metadata in PostgreSQL
            knowledge_doc = KnowledgeDocument(
                document_id=doc_id,
                customer_id=int(customer_id),
                filename=filename,
                content_type=content_type,
                file_size=len(file_content),
                runpod_storage_path=storage_path,
                runpod_url=file_url,
                text_content=extracted_text[:50000],  # Store first 50k chars
                text_length=len(extracted_text),
                embedding=embedding,
                icp_data=icp_data,
                processing_status='completed'
            )

            db.add(knowledge_doc)
            db.commit()
            db.refresh(knowledge_doc)

            logger.info(f"Uploaded document {doc_id} for customer {customer_id}")

            return {
                'document_id': doc_id,
                'filename': filename,
                'file_url': file_url,
                'text_preview': extracted_text[:500],
                'text_length': len(extracted_text),
                'embedding_dimension': self.embedding_dimension,
                'icp_criteria': icp_data,
                'created_at': knowledge_doc.created_at.isoformat()
            }

        except Exception as e:
            logger.error(f"Failed to upload document for customer {customer_id}: {e}", exc_info=True)
            db.rollback()
            raise

    def _extract_text(self, file_content: bytes, filename: str, content_type: str) -> str:
        """
        Extract text from various document formats

        Args:
            file_content: Document bytes
            filename: Original filename
            content_type: MIME type

        Returns:
            Extracted text content
        """
        try:
            # Determine format from content type or filename
            if 'pdf' in content_type.lower() or filename.lower().endswith('.pdf'):
                return self._extract_from_pdf(file_content)

            elif 'word' in content_type.lower() or filename.lower().endswith(('.docx', '.doc')):
                return self._extract_from_docx(file_content)

            elif 'text' in content_type.lower() or filename.lower().endswith('.txt'):
                return file_content.decode('utf-8')

            else:
                # Try UTF-8 decode as fallback
                return file_content.decode('utf-8', errors='ignore')

        except Exception as e:
            logger.error(f"Failed to extract text from {filename}: {e}")
            raise ValueError(f"Unsupported document format or corrupted file: {filename}")

    def _extract_from_pdf(self, file_content: bytes) -> str:
        """Extract text from PDF file"""
        try:
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            text_parts = []

            for page in pdf_reader.pages:
                text_parts.append(page.extract_text())

            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"PDF extraction error: {e}")
            raise

    def _extract_from_docx(self, file_content: bytes) -> str:
        """Extract text from DOCX file"""
        try:
            doc = DocxDocument(io.BytesIO(file_content))
            text_parts = [paragraph.text for paragraph in doc.paragraphs if paragraph.text.strip()]
            return "\n\n".join(text_parts)
        except Exception as e:
            logger.error(f"DOCX extraction error: {e}")
            raise

    def _generate_embedding(self, text: str) -> List[float]:
        """
        Generate vector embedding for text

        Args:
            text: Input text

        Returns:
            384-dimensional embedding vector
        """
        try:
            # Truncate text if too long (model max is typically 512 tokens)
            max_length = 8000  # characters, roughly ~2000 tokens
            if len(text) > max_length:
                text = text[:max_length]

            # Generate embedding
            embedding = self.embedding_model.encode(text, show_progress_bar=False)

            # Convert to list for JSON serialization
            return embedding.tolist()

        except Exception as e:
            logger.error(f"Failed to generate embedding: {e}")
            raise

    def _extract_icp_criteria(self, text: str) -> Dict:
        """
        Extract ICP (Ideal Customer Profile) criteria from document text

        Uses keyword matching and pattern recognition to identify:
        - Target industries
        - Company size ranges
        - Geographic regions
        - Key decision maker titles

        Args:
            text: Document text

        Returns:
            Dictionary of ICP criteria
        """
        text_lower = text.lower()

        # Common industry keywords
        industries = []
        industry_keywords = {
            'saas': ['saas', 'software as a service', 'cloud software'],
            'fintech': ['fintech', 'financial technology', 'banking', 'payments'],
            'healthcare': ['healthcare', 'health tech', 'medical', 'hospital'],
            'ecommerce': ['ecommerce', 'e-commerce', 'retail', 'online store'],
            'manufacturing': ['manufacturing', 'industrial', 'factory'],
            'logistics': ['logistics', 'supply chain', 'transportation'],
            'education': ['education', 'edtech', 'learning', 'university'],
            'real estate': ['real estate', 'property', 'construction']
        }

        for industry, keywords in industry_keywords.items():
            if any(kw in text_lower for kw in keywords):
                industries.append(industry)

        # Company size indicators
        company_sizes = []
        size_patterns = {
            'enterprise': ['enterprise', '1000+ employees', 'large corporation'],
            'mid-market': ['mid-market', '100-1000 employees', 'medium business'],
            'smb': ['small business', 'smb', '10-100 employees', 'startup']
        }

        for size_category, patterns in size_patterns.items():
            if any(pattern in text_lower for pattern in patterns):
                company_sizes.append(size_category)

        # Decision maker titles
        decision_makers = []
        title_keywords = ['ceo', 'cto', 'cfo', 'vp', 'director', 'head of', 'chief']

        for keyword in title_keywords:
            if keyword in text_lower:
                decision_makers.append(keyword.upper() if len(keyword) <= 3 else keyword.title())

        # Geographic regions
        regions = []
        region_keywords = {
            'North America': ['usa', 'united states', 'canada', 'north america'],
            'Europe': ['europe', 'eu', 'uk', 'germany', 'france'],
            'Asia': ['asia', 'china', 'japan', 'india', 'singapore'],
            'Global': ['global', 'worldwide', 'international']
        }

        for region, keywords in region_keywords.items():
            if any(kw in text_lower for kw in keywords):
                regions.append(region)

        return {
            'target_industries': list(set(industries)),
            'company_sizes': list(set(company_sizes)),
            'decision_makers': list(set(decision_makers)),
            'target_regions': list(set(regions)),
            'extracted_at': datetime.now().isoformat()
        }

    def search_similar_documents(
        self,
        query_text: str,
        customer_id: str,
        limit: int = 10,
        similarity_threshold: float = 0.7,
        db: Session = None
    ) -> List[Dict]:
        """
        Search for similar documents using vector similarity (pgvector)

        Args:
            query_text: Search query
            customer_id: Customer ID for filtering
            limit: Maximum results
            similarity_threshold: Minimum cosine similarity (0-1)
            db: Database session

        Returns:
            List of similar documents with similarity scores
        """
        try:
            # Generate query embedding
            query_embedding = self._generate_embedding(query_text)

            # Query PostgreSQL for customer documents
            # Using pgvector for similarity search
            documents = db.query(KnowledgeDocument).filter(
                KnowledgeDocument.customer_id == int(customer_id),
                KnowledgeDocument.processing_status == 'completed'
            ).limit(limit).all()

            if not documents:
                return []

            # For now, return documents without vector similarity scoring
            # TODO: Implement pgvector cosine similarity once pgvector extension is enabled
            results = []
            for doc in documents:
                results.append({
                    'document_id': doc.document_id,
                    'filename': doc.filename,
                    'file_url': doc.runpod_url,
                    'icp_criteria': doc.icp_data,
                    'created_at': doc.created_at.isoformat(),
                    'similarity_score': 0.85  # Placeholder until pgvector integration
                })

            logger.info(f"Found {len(results)} documents for customer {customer_id}")
            return results

        except Exception as e:
            logger.error(f"Failed to search documents for customer {customer_id}: {e}")
            raise

    def get_customer_documents(
        self,
        customer_id: str,
        limit: Optional[int] = 50,
        db: Session = None
    ) -> List[Dict]:
        """
        Get all documents for a customer

        Args:
            customer_id: Customer ID
            limit: Maximum results
            db: Database session

        Returns:
            List of document metadata
        """
        try:
            documents = db.query(KnowledgeDocument).filter(
                KnowledgeDocument.customer_id == int(customer_id)
            ).order_by(
                KnowledgeDocument.created_at.desc()
            ).limit(limit).all()

            result = []
            for doc in documents:
                result.append({
                    'document_id': doc.document_id,
                    'customer_id': doc.customer_id,
                    'filename': doc.filename,
                    'content_type': doc.content_type,
                    'file_url': doc.runpod_url,
                    'text_length': doc.text_length,
                    'icp_criteria': doc.icp_data,
                    'processing_status': doc.processing_status,
                    'created_at': doc.created_at.isoformat()
                })

            logger.info(f"Retrieved {len(result)} documents for customer {customer_id}")
            return result

        except Exception as e:
            logger.error(f"Failed to get documents for customer {customer_id}: {e}")
            raise

    def delete_document(self, document_id: str, customer_id: str, db: Session):
        """
        Delete a customer document

        Args:
            document_id: Document ID
            customer_id: Customer ID (for authorization)
            db: Database session
        """
        try:
            # Get document metadata from PostgreSQL
            doc = db.query(KnowledgeDocument).filter(
                KnowledgeDocument.document_id == document_id
            ).first()

            if not doc:
                raise ValueError(f"Document {document_id} not found")

            # Verify customer ownership
            if str(doc.customer_id) != str(customer_id):
                raise PermissionError(f"Customer {customer_id} does not own document {document_id}")

            # Delete from RunPod S3 Storage
            storage_path = doc.runpod_storage_path
            if storage_path:
                try:
                    self.storage.delete_file(storage_path)
                except Exception as e:
                    logger.warning(f"Failed to delete file from RunPod: {e}")

            # Delete from PostgreSQL
            db.delete(doc)
            db.commit()

            logger.info(f"Deleted document {document_id} for customer {customer_id}")

        except Exception as e:
            logger.error(f"Failed to delete document {document_id}: {e}")
            db.rollback()
            raise
