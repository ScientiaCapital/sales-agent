"""
Document Processing Service

AI-powered analysis of PDF, DOCX, and TXT documents for lead enrichment.
Supports resume parsing, job matching, and skill extraction.
"""

import io
import logging
import re
from typing import Dict, Any, List, Optional, Tuple
from datetime import datetime

# PDF processing - using pdfplumber (best for text extraction per Context7)
try:
    import pdfplumber
except ImportError:
    pdfplumber = None

# DOCX processing
try:
    import docx
except ImportError:
    docx = None

from fastapi import HTTPException
from app.services.cerebras import CerebrasService

logger = logging.getLogger(__name__)


class DocumentProcessor:
    """
    Service for processing and analyzing documents with AI

    Supported formats:
    - PDF (.pdf) - Using pdfplumber for best text extraction
    - DOCX (.docx) - Microsoft Word documents
    - TXT (.txt) - Plain text files

    AI capabilities:
    - Resume parsing (skills, experience, education)
    - Job matching and fit scoring
    - Contact information extraction
    - Qualification assessment
    """

    def __init__(self):
        self.cerebras = CerebrasService()
        
        # Check dependencies
        if not pdfplumber:
            logger.warning("pdfplumber not installed - PDF processing disabled")
        if not docx:
            logger.warning("python-docx not installed - DOCX processing disabled")

    def extract_text_from_pdf(self, file_content: bytes) -> str:
        """
        Extract text from PDF using pdfplumber

        Args:
            file_content: PDF file bytes

        Returns:
            Extracted text content

        Raises:
            HTTPException: If PDF processing fails
        """
        if not pdfplumber:
            raise HTTPException(
                status_code=501,
                detail="PDF processing not available - pdfplumber not installed"
            )

        try:
            pdf_file = io.BytesIO(file_content)
            text_parts = []

            with pdfplumber.open(pdf_file) as pdf:
                for page_num, page in enumerate(pdf.pages, start=1):
                    # Extract text with whitespace preservation
                    page_text = page.extract_text(keep_blank_chars=True)
                    if page_text:
                        text_parts.append(f"=== Page {page_num} ===\n{page_text}")

                    # Also extract tables if present
                    tables = page.extract_tables()
                    for table_num, table in enumerate(tables, start=1):
                        if table:
                            table_text = "\n".join(["\t".join(str(cell) if cell else "" for cell in row) for row in table])
                            text_parts.append(f"=== Table {page_num}.{table_num} ===\n{table_text}")

            full_text = "\n\n".join(text_parts)
            
            if not full_text.strip():
                raise HTTPException(
                    status_code=400,
                    detail="PDF appears to be empty or contains only images"
                )

            return full_text

        except Exception as e:
            logger.error(f"PDF extraction failed: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=400,
                detail=f"Failed to extract text from PDF: {str(e)}"
            )

    def extract_text_from_docx(self, file_content: bytes) -> str:
        """
        Extract text from DOCX file

        Args:
            file_content: DOCX file bytes

        Returns:
            Extracted text content

        Raises:
            HTTPException: If DOCX processing fails
        """
        if not docx:
            raise HTTPException(
                status_code=501,
                detail="DOCX processing not available - python-docx not installed"
            )

        try:
            docx_file = io.BytesIO(file_content)
            doc = docx.Document(docx_file)

            text_parts = []

            # Extract paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            # Extract tables
            for table_num, table in enumerate(doc.tables, start=1):
                table_text = f"=== Table {table_num} ==="
                for row in table.rows:
                    row_text = "\t".join([cell.text for cell in row.cells])
                    table_text += f"\n{row_text}"
                text_parts.append(table_text)

            full_text = "\n\n".join(text_parts)

            if not full_text.strip():
                raise HTTPException(
                    status_code=400,
                    detail="DOCX document appears to be empty"
                )

            return full_text

        except Exception as e:
            logger.error(f"DOCX extraction failed: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=400,
                detail=f"Failed to extract text from DOCX: {str(e)}"
            )

    def extract_text_from_txt(self, file_content: bytes) -> str:
        """
        Extract text from TXT file

        Args:
            file_content: TXT file bytes

        Returns:
            Extracted text content

        Raises:
            HTTPException: If text decoding fails
        """
        try:
            text = file_content.decode('utf-8')
            
            if not text.strip():
                raise HTTPException(
                    status_code=400,
                    detail="Text file is empty"
                )

            return text

        except UnicodeDecodeError:
            # Try other encodings
            for encoding in ['latin-1', 'cp1252', 'iso-8859-1']:
                try:
                    text = file_content.decode(encoding)
                    return text
                except UnicodeDecodeError:
                    continue

            raise HTTPException(
                status_code=400,
                detail="Unable to decode text file - unsupported encoding"
            )

    def extract_text(self, filename: str, file_content: bytes) -> str:
        """
        Extract text from document based on file extension

        Args:
            filename: Original filename with extension
            file_content: File bytes

        Returns:
            Extracted text content

        Raises:
            HTTPException: If file type unsupported or extraction fails
        """
        ext = filename.lower().split('.')[-1]

        if ext == 'pdf':
            return self.extract_text_from_pdf(file_content)
        elif ext in ['docx', 'doc']:
            return self.extract_text_from_docx(file_content)
        elif ext == 'txt':
            return self.extract_text_from_txt(file_content)
        else:
            raise HTTPException(
                status_code=400,
                detail=f"Unsupported file type: .{ext}. Supported: PDF, DOCX, TXT"
            )

    def analyze_resume(
        self,
        text: str,
        job_description: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Analyze resume with AI to extract skills, experience, and job fit

        Args:
            text: Resume text content
            job_description: Optional job description for matching

        Returns:
            Dictionary with analysis results
        """
        # Build AI prompt
        system_prompt = """You are an expert HR assistant analyzing resumes.
Extract and structure information from the resume in this exact JSON format:

{
    "contact_info": {
        "name": "Full name",
        "email": "email@example.com",
        "phone": "+1234567890",
        "linkedin": "linkedin.com/in/profile"
    },
    "skills": ["skill1", "skill2", "skill3"],
    "experience_years": <number>,
    "education": [
        {"degree": "Degree name", "institution": "School", "year": 2020}
    ],
    "work_experience": [
        {"title": "Job Title", "company": "Company", "years": 2}
    ],
    "key_strengths": ["strength1", "strength2"],
    "summary": "2-3 sentence professional summary"
}

If job description provided, also include:
    "job_fit_score": <0-100>,
    "job_fit_reasoning": "Why this candidate matches/doesn't match"
"""

        user_prompt = f"Resume:\n\n{text[:4000]}"  # Limit to prevent token overflow

        if job_description:
            user_prompt += f"\n\nJob Description:\n\n{job_description[:1000]}"

        try:
            # Call Cerebras for analysis
            score, reasoning, latency_ms = self.cerebras.qualify_lead(
                company_name="Resume Analysis",
                notes=user_prompt
            )

            # Parse AI response (simplified - in production use structured output)
            import json
            
            # Attempt to extract JSON from response
            try:
                # Try direct JSON parse
                analysis = json.loads(reasoning)
            except json.JSONDecodeError:
                # Fallback: basic extraction
                analysis = {
                    "summary": reasoning,
                    "ai_score": score,
                    "processing_latency_ms": latency_ms
                }

            return analysis

        except Exception as e:
            logger.error(f"Resume analysis failed: {str(e)}", exc_info=True)
            raise HTTPException(
                status_code=500,
                detail=f"Resume analysis failed: {str(e)}"
            )

    def process_document(
        self,
        filename: str,
        file_content: bytes,
        document_type: str = "resume",
        job_description: Optional[str] = None
    ) -> Dict[str, Any]:
        """
        Process document end-to-end: extract text and analyze

        Args:
            filename: Original filename
            file_content: File bytes
            document_type: Type of document (resume, contract, etc.)
            job_description: Optional job description for matching

        Returns:
            Complete analysis results
        """
        start_time = datetime.now()

        # Extract text
        text = self.extract_text(filename, file_content)
        logger.info(f"Extracted {len(text)} characters from {filename}")

        # Analyze based on document type
        if document_type == "resume":
            analysis = self.analyze_resume(text, job_description)
        else:
            # Generic document analysis
            analysis = {
                "text_length": len(text),
                "word_count": len(text.split()),
                "preview": text[:500]
            }

        end_time = datetime.now()
        duration_ms = int((end_time - start_time).total_seconds() * 1000)

        return {
            "filename": filename,
            "document_type": document_type,
            "processing_duration_ms": duration_ms,
            "text_length": len(text),
            "analysis": analysis
        }
